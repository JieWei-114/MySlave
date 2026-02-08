"""
File Extraction Service

Extracts text content from various file formats.
Supports PDF, DOCX, and plain text files.

Supported Formats:
- .pdf: PyPDF2 library
- .docx: python-docx library
- .doc: Not supported (requires conversion to .docx)
- Text files: Direct UTF-8 decoding

Features:
- Automatic format detection based on file extension
- Content truncation to prevent memory issues
- Error handling with user-friendly messages
- Encoding fallback (UTF-8 -> Latin-1)
- File attachment persistence to MongoDB
"""
import io
import logging
import uuid
from datetime import datetime, timedelta

from app.config.settings import settings
from app.core.db import file_attachments_collection

logger = logging.getLogger(__name__)

# Centralized file upload limit
FILE_UPLOAD_MAX_CHARS = settings.FILE_UPLOAD_MAX_CHARS


def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """
    Extract text content from various file types.
    
    Automatically detects file type based on extension and routes to appropriate
    extraction function. Applies content truncation based on settings.
    
    """
    filename_lower = filename.lower()

    logger.info(f'Extracting content from {filename} ({len(file_content)} bytes)')

    try:
        if filename_lower.endswith('.pdf'):
            content = extract_from_pdf(file_content)
        elif filename_lower.endswith('.docx'):
            content = extract_from_docx(file_content)
        elif filename_lower.endswith('.doc'):
            raise ValueError(
                'Legacy .doc format not supported. '
                'Please convert the file to .docx format using Microsoft Word or LibreOffice'
            )
        else:
            content = extract_as_text(file_content)

        # Truncate using centralized limit
        content = truncate_content(content)

        logger.info(f'Successfully extracted {len(content)} chars from {filename}')
        return content

    except Exception as e:
        logger.error(f'Failed to extract content from {filename}: {e}')
        raise ValueError(f'Failed to extract content from {filename}: {str(e)}')


def extract_from_pdf(file_content: bytes) -> str:
    """
    Extract text from PDF file using PyPDF2.
    
    Reads all pages and concatenates text content.
    Handles various PDF formats and encodings.
    
    """
    try:
        import PyPDF2

        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)

        text_parts = []
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                text = page.extract_text()
                if text.strip():
                    text_parts.append(f'--- Page {page_num + 1} ---\n{text}')
            except Exception as e:
                logger.warning(f'Failed to extract page {page_num + 1}: {e}')
                continue

        if not text_parts:
            raise ValueError('No text content found in PDF')

        return '\n\n'.join(text_parts)

    except ImportError:
        raise ValueError('PDF extraction not available. Install PyPDF2: pip install PyPDF2')


def extract_from_docx(file_content: bytes) -> str:
    """
    Extract text from .docx file.
    """
    try:
        from docx import Document

        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)

        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        if not text_parts:
            raise ValueError('No text content found in Word document')

        return '\n\n'.join(text_parts)

    except ImportError:
        raise ValueError(
            'Word extraction not available. Install python-docx: pip install python-docx'
        )


def extract_as_text(file_content: bytes) -> str:
    """
    Extract content as plain text with encoding detection.
    """
    # Try common encodings
    encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

    for encoding in encodings:
        try:
            text = file_content.decode(encoding)
            if text.strip():
                return text.strip()
        except (UnicodeDecodeError, AttributeError):
            continue

    raise ValueError('Unable to decode file content as text')


def truncate_content(content: str, max_chars: int = None) -> str:
    """
    Truncate content to maximum length.
    """
    if max_chars is None:
        max_chars = FILE_UPLOAD_MAX_CHARS

    if len(content) <= max_chars:
        return content

    truncated = content[:max_chars]
    return f'{truncated}\n\n[Content truncated to {max_chars} chars - original length: {len(content)} chars]'


# ============================================================
# FILE ATTACHMENT PERSISTENCE
# ============================================================

def store_file_attachment(
    session_id: str,
    filename: str,
    content: str,
    file_type: str,
) -> dict:
    # Store an uploaded file attachment to database
    if len(content) > settings.FILE_ATTACHMENT_MAX_CHARS:
        content = content[:settings.FILE_ATTACHMENT_MAX_CHARS]
        logger.warning(f"File {filename} truncated to {settings.FILE_ATTACHMENT_MAX_CHARS} chars")

    expires_at = datetime.utcnow() + timedelta(days=settings.FILE_ATTACHMENT_EXPIRY_DAYS)

    file_attachment = {
        "id": str(uuid.uuid4()),
        "session_id": session_id,
        "filename": filename,
        "file_type": file_type,
        "size_bytes": len(content.encode()),
        "size_chars": len(content),
        "content": content,
        "uploaded_at": datetime.utcnow(),
        "expires_at": expires_at,
    }

    result = file_attachments_collection.insert_one(file_attachment)
    file_attachment["_id"] = result.inserted_id
    logger.info(f"File attachment stored: {filename} (ID: {file_attachment['id']})")
    return file_attachment


def get_file_attachment(file_id: str) -> dict | None:
    # Retrieve a file attachment
    return file_attachments_collection.find_one({"id": file_id}, {"_id": 0})


def list_file_attachments(session_id: str) -> list[dict]:
    # List file attachments for a session (without content)
    cursor = file_attachments_collection.find(
        {"session_id": session_id},
        {"_id": 0, "content": 0},
    ).sort("uploaded_at", -1)
    return list(cursor)


def delete_file_attachment_for_session(session_id: str, file_id: str) -> bool:
    # Delete a file attachment for a specific session
    result = file_attachments_collection.delete_one({"id": file_id, "session_id": session_id})
    return result.deleted_count > 0


def delete_file_attachments_for_session(session_id: str) -> int:
    # Delete all file attachments for a session
    result = file_attachments_collection.delete_many({"session_id": session_id})
    logger.info(f"Deleted {result.deleted_count} file attachments for session {session_id}")
    return result.deleted_count